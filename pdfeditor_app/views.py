import os, json, logging
from django.shortcuts import render, redirect
from django.http import HttpResponse
from .models import Document
from django.shortcuts import get_object_or_404
from .forms import UploadFileForm
import pdfplumber
from docx import Document as DocxDocument
from reportlab.lib.pagesizes import A4
from reportlab.pdfgen import canvas
import os

logger = logging.getLogger(__name__)

def upload_pdf(request):
    if request.method == 'POST':
        form = UploadFileForm(request.POST, request.FILES)
        if form.is_valid():
            doc = Document(pdf_file=request.FILES['pdf_file'])
            doc.save()

            # Extract text + tables
            full_text = ""
            tables = []
            with pdfplumber.open(doc.pdf_file.path) as pdf:
                for page in pdf.pages:
                    full_text += page.extract_text() or ""
                    for table in page.extract_tables():
                        headers = table[0]
                        rows = table[1:]
                        table_dict = {headers[i]: [row[i] for row in rows] for i in range(len(headers))}
                        tables.append(table_dict)

            doc.full_text = full_text
            doc.tables_json = tables
            doc.save()

            return redirect('edit_document', doc.id)
    else:
        form = UploadFileForm()
    return render(request, 'upload.html', {'form': form})


def edit_document(request, doc_id):
    doc = Document.objects.get(id=doc_id)

    if request.method == 'POST':
        doc.full_text = request.POST.get("full_text", "")
        doc.tables_json = json.loads(request.POST.get("tables_json", "[]"))
        doc.save()

        # Generate Word file
        word_path = f"media/outputs/doc_{doc.id}.docx"
        word_doc = DocxDocument()
        word_doc.add_paragraph(doc.full_text)
        for tbl in doc.tables_json:
            headers = list(tbl.keys())
            rows = list(zip(*tbl.values()))
            table = word_doc.add_table(rows=1, cols=len(headers))
            hdr_cells = table.rows[0].cells
            for i, h in enumerate(headers):
                hdr_cells[i].text = h
            for row in rows:
                row_cells = table.add_row().cells
                for i, val in enumerate(row):
                    row_cells[i].text = str(val)
        word_doc.save(word_path)
        doc.word_file.name = word_path.replace("media/", "")
        doc.save()

        # Generate PDF file
        pdf_path = f"media/outputs/doc_{doc.id}.pdf"
        c = canvas.Canvas(pdf_path, pagesize=A4)
        text_obj = c.beginText(50, 800)
        text_obj.setFont("Helvetica", 10)
        for line in doc.full_text.split("\n"):
            text_obj.textLine(line)
        c.drawText(text_obj)

        y = 600
        for tbl in doc.tables_json:
            headers = list(tbl.keys())
            rows = list(zip(*tbl.values()))
            c.drawString(50, y, " | ".join(headers))
            y -= 20
            for row in rows:
                c.drawString(50, y, " | ".join([str(v) for v in row]))
                y -= 20
            y -= 20
        c.save()

        doc.pdf_output_file.name = pdf_path.replace("media/", "")
        doc.save()

        return redirect('edit_document', doc.id)

    return render(request, 'edit.html', {"doc": doc, "tables_json": json.dumps(doc.tables_json)})


def download_document(request, doc_id):
    doc = get_object_or_404(Document, id=doc_id)
    file_type = request.GET.get("type", "pdf")  # default = pdf

    if file_type == "pdf" and doc.pdf_output_file:
        file_field = doc.pdf_output_file
        content_type = "application/pdf"
        filename = "converted.pdf"
    elif file_type == "word" and doc.word_file:
        file_field = doc.word_file
        content_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        filename = "converted.docx"
    else:
        return HttpResponse("Requested file not available", status=404)

    file_path = file_field.path
    with open(file_path, "rb") as f:
        response = HttpResponse(f.read(), content_type=content_type)
        response["Content-Disposition"] = f'attachment; filename="{filename}"'
        return response